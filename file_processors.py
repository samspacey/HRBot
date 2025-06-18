"""
file_processors.py - Multi-format document processors for interactive chatbot.

This module provides document loaders and processors for different file formats
including PDF, TXT, DOCX, and Markdown files.
"""

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from langchain.schema import Document
from langchain_community.document_loaders import PyPDFLoader
from validation import get_validator, ValidationError

logger = logging.getLogger(__name__)

class MultiFormatDocumentProcessor:
    """
    Processes documents of various formats into LangChain Document objects.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.validator = get_validator()
        self.processors = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.docx': self._process_docx,
        }
    
    def process_file(self, file_path: str, metadata: Optional[Dict] = None) -> List[Document]:
        """
        Process a single file and return Document objects.
        
        Args:
            file_path: Path to the file to process
            metadata: Optional metadata to include with documents
            
        Returns:
            List of Document objects
            
        Raises:
            ValidationError: If file validation fails
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        try:
            # Validate file path
            validated_path = self.validator.validate_file_path(file_path)
            file_path = str(validated_path)
            
            # Get file extension
            file_ext = Path(file_path).suffix.lower()
            
            if file_ext not in self.processors:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Process file with appropriate processor
            processor = self.processors[file_ext]
            documents = processor(file_path)
            
            # Add metadata to all documents
            if metadata:
                for doc in documents:
                    doc.metadata.update(metadata)
            
            # Add file-specific metadata
            file_metadata = {
                'source': file_path,
                'filename': Path(file_path).name,
                'file_type': file_ext,
                'file_size': os.path.getsize(file_path)
            }
            
            for doc in documents:
                doc.metadata.update(file_metadata)
            
            logger.info(f"Processed {file_ext} file: {file_path} -> {len(documents)} documents")
            return documents
            
        except Exception as e:
            logger.error(f"Error processing file {file_path}: {str(e)}")
            raise
    
    def process_files(self, file_paths: List[str], metadata: Optional[Dict] = None) -> List[Document]:
        """
        Process multiple files and return combined Document objects.
        
        Args:
            file_paths: List of file paths to process
            metadata: Optional metadata to include with all documents
            
        Returns:
            List of Document objects from all files
        """
        all_documents = []
        errors = []
        
        for file_path in file_paths:
            try:
                documents = self.process_file(file_path, metadata)
                all_documents.extend(documents)
            except Exception as e:
                error_msg = f"Failed to process {file_path}: {str(e)}"
                errors.append(error_msg)
                logger.error(error_msg)
        
        if errors:
            logger.warning(f"Processing completed with {len(errors)} errors: {'; '.join(errors)}")
        
        logger.info(f"Processed {len(file_paths)} files -> {len(all_documents)} total documents")
        return all_documents
    
    def _process_pdf(self, file_path: str) -> List[Document]:
        """Process PDF file using PyPDFLoader."""
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            
            # Add page numbers to metadata
            for i, doc in enumerate(documents):
                doc.metadata['page'] = i + 1
            
            return documents
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise ValueError(f"Failed to process PDF file: {str(e)}")
    
    def _process_text(self, file_path: str) -> List[Document]:
        """Process plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Create single document for text file
            document = Document(
                page_content=content,
                metadata={'page': 1}
            )
            
            return [document]
            
        except UnicodeDecodeError:
            # Try different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    document = Document(
                        page_content=content,
                        metadata={'page': 1, 'encoding': encoding}
                    )
                    
                    logger.warning(f"Used {encoding} encoding for {file_path}")
                    return [document]
                    
                except UnicodeDecodeError:
                    continue
            
            raise ValueError(f"Could not decode text file {file_path} with any supported encoding")
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise ValueError(f"Failed to process text file: {str(e)}")
    
    def _process_markdown(self, file_path: str) -> List[Document]:
        """Process Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # For now, treat as plain text
            # Could be enhanced to parse markdown structure
            document = Document(
                page_content=content,
                metadata={'page': 1, 'format': 'markdown'}
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"Error processing markdown file {file_path}: {str(e)}")
            raise ValueError(f"Failed to process markdown file: {str(e)}")
    
    def _process_docx(self, file_path: str) -> List[Document]:
        """Process DOCX file."""
        try:
            # Import docx here to handle optional dependency
            try:
                from docx import Document as DocxDocument
            except ImportError:
                raise ValueError("python-docx is required to process DOCX files. Install with: pip install python-docx")
            
            # Load DOCX document
            docx_doc = DocxDocument(file_path)
            
            # Extract text from all paragraphs
            paragraphs = []
            for paragraph in docx_doc.paragraphs:
                if paragraph.text.strip():  # Skip empty paragraphs
                    paragraphs.append(paragraph.text)
            
            # Combine paragraphs into single content
            content = '\n\n'.join(paragraphs)
            
            if not content.strip():
                raise ValueError("DOCX file appears to be empty or contains no readable text")
            
            # Create single document
            document = Document(
                page_content=content,
                metadata={'page': 1, 'format': 'docx', 'paragraphs': len(paragraphs)}
            )
            
            return [document]
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.processors.keys())
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file format is supported."""
        file_ext = Path(file_path).suffix.lower()
        return file_ext in self.processors
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get information about a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file information
        """
        try:
            path = Path(file_path)
            
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_info = {
                'name': path.name,
                'stem': path.stem,
                'extension': path.suffix.lower(),
                'size': path.stat().st_size,
                'size_mb': round(path.stat().st_size / (1024 * 1024), 2),
                'supported': self.is_supported(file_path),
                'absolute_path': str(path.absolute())
            }
            
            return file_info
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {str(e)}")
            raise


def get_document_processor() -> MultiFormatDocumentProcessor:
    """
    Get document processor instance.
    
    Returns:
        MultiFormatDocumentProcessor instance
    """
    return MultiFormatDocumentProcessor()


def process_uploaded_files(file_paths: List[str], metadata: Optional[Dict] = None) -> List[Document]:
    """
    Convenience function to process multiple uploaded files.
    
    Args:
        file_paths: List of file paths to process
        metadata: Optional metadata to include with documents
        
    Returns:
        List of Document objects
    """
    processor = get_document_processor()
    return processor.process_files(file_paths, metadata)